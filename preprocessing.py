import pandas as pd
import re
from generate_text import summarize
from resources import questions_data, empty_collector, levels, full_names, report_data
import copy
import asyncio
from docx import Document
from docx.shared import Inches
import plotly.graph_objects as go


class Preprocessing:
    def __init__(self, path):
        self.poll_data = pd.read_excel(path).dropna(subset=[questions_data[1]["column_name"], questions_data[9]["column_name"], questions_data[16]["column_name"]])

    def collect(self):
        def get_initials(arr):
            return [word[0] for word in arr]

        def record_respond():
            nonlocal level, name, collector, row
            for question in levels[level]["relevant_questions"]:
                if questions_data[question]["type"] == "number":
                    try:
                        collector[name][question]["value"] += int(row[questions_data[question]["column_name"]])
                    except ValueError:
                        continue
                    else:
                        collector[name][question]["quantity"] += 1
                elif questions_data[question]["type"] == "select":
                    try:
                        collector[name][question]["value"] += questions_data[question]["estimate"][str(row[questions_data[question]["column_name"]])]
                    except ValueError:
                        continue
                    except KeyError:
                        continue
                    else:
                        collector[name][question]["quantity"] += 1
                else:
                    try:
                        free_respond = str(row[questions_data[question]["column_name"]])
                    except ValueError:
                        continue
                    else:
                        if len(re.findall('[а-я]+', free_respond)) > 0:
                            collector[name][question].append(free_respond)

        collector = {}
        for index, row in self.poll_data.iterrows():
            for level in levels.keys():
                name = re.findall('[а-я]+', row[questions_data[levels[level]["full_name"]]["column_name"]].lower())
                if len(name) > 3 or len(name) == 0:
                    continue

                initials = []
                for i in range(len(name) - 1, -1, -1):
                    if len(name[i]) == 1:
                        initials.append(name[i])
                        name.pop(i)

                intersect = []
                name = set(name)

                if len(initials) > 0:
                    initials.sort()
                    for full_name in full_names.keys():
                        if full_name.issuperset(name):
                            fn_initials = sorted(get_initials(set(full_name) - name))
                            counter = 0
                            for i in range(len(fn_initials)):
                                if fn_initials[i] == initials[counter]:
                                    counter += 1
                                    if counter == len(initials):
                                        intersect.append(full_name)
                                        break
                else:
                    for full_name in full_names.keys():
                        if full_name.issuperset(name):
                            intersect.append(full_name)

                if len(intersect) == 1:
                    name = intersect[0]
                else:
                    name = frozenset(name)

                if name not in collector:
                    collector[name] = copy.deepcopy(empty_collector)
                record_respond()

        return self.free_responses(self.numeric_responses(collector))



    def free_responses(self, collector):
        def to_summarize(name, question):
            if len(collector[name][question]) > 0:
                reviews = ""
                for i in range(len(collector[name][question])):
                    reviews += f"{i + 1}: {collector[name][question][i]}\n"
                return asyncio.run(summarize(reviews))

        for name in collector.keys():
            collector[name][8] = to_summarize(name, 8)
            collector[name][14] = to_summarize(name, 14)
            collector[name][22] = to_summarize(name, 22)

        return collector

    def numeric_responses(self, collector):
        for name in collector.keys():
            for question in collector[name].keys():
                if type(collector[name][question]) is dict:
                    if collector[name][question]["quantity"] < 1:
                        collector[name][question] = None
                    else:
                        collector[name][question] = round(collector[name][question]["value"] / collector[name][question]["quantity"], 1)

        return collector

    def to_csv(self, collector):
        return pd.DataFrame.from_dict(collector, orient='index')

    def get_mean(self, collector):
        counter = 0
        s = 0
        for name in collector.keys():
            if collector[name][4] is not None:
                s += collector[name][4]
                counter += 1
            if collector[name][11] is not None:
                s += collector[name][11]
                counter += 1
            if collector[name][20] is not None:
                s += collector[name][20]
                counter += 1
        return round(s / counter, 1)

    def to_report(self, collector):
        for name in collector.keys():
            document = Document()

            if name in full_names:
                fn = full_names[name]
            else:
                fn = ' '.join(map(lambda x: x.capitalize(), list(name)))

            document.add_heading(fn, 0)

            document.add_paragraph(f"Cредний урочень оценок: {self.get_mean(collector)}")
            for level in levels:
                graph = {"x": [], "y": []}
                document.add_paragraph(f"Уровень подчинения: {level}", style='Intense Quote')
                for question in sorted(list(levels[level]["relevant_questions"])):
                    if questions_data[question]["type"] == "number":
                        if collector[name][question] is not None:
                            document.add_paragraph(f"Средняя оценка по уровню: {collector[name][question]}")
                        else:
                            document.add_paragraph("Недостаточно оценок")
                    elif questions_data[question]["type"] == "select":
                        if collector[name][question] is not None:
                            graph["x"].append(collector[name][question])
                            graph["y"].append(report_data[question])
                    else:
                        if collector[name][question] is not None:
                            document.add_paragraph(f"Cуммаризированные отзывы с свободным ответом: {collector[name][question]}")
                        else:
                            document.add_paragraph("Недостаточно отзывов с свободным ответом")

                fig = go.Figure(go.Bar(
                    x=graph["x"],
                    y=graph["y"],
                    orientation='h'))
                fig.write_image("fig.png")
                document.add_picture("fig.png", width=Inches(8))

                document.add_page_break()
            document.save(f'{fn}.docx')


if __name__ == '__main__':
    path = input("Please, enter path to xlsx file:")
    prep = Preprocessing(path)
    collector = prep.collect()
    prep.to_report(collector)
    print(collector)


